# -*- coding: utf-8 -*-
"""
Word Document Styles Module

Single source of truth for DOCX formatting.
All document generators must import from this module.

IMPORTANT: This module uses Word's built-in Styles system.
You can easily modify formatting by:
1. Opening a generated document in Word
2. Right-click any styled element → Modify Style
3. Changes apply to all elements using that style

See docx-formatting-spec.md for the complete specification.
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ROW_HEIGHT_RULE
from docx.enum.section import WD_ORIENT
from docx.oxml import parse_xml
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# =============================================================================
# CONSTANTS - Based on docx-formatting-spec.md
# =============================================================================

# Fonts
FONT_PRIMARY = 'Aptos'
FONT_DISPLAY = 'Aptos Display'

# Colours (RGB tuples)
COLOUR_BLACK = RGBColor(0, 0, 0)
COLOUR_GREY = RGBColor(128, 128, 128)
COLOUR_DARK_GREY = RGBColor(64, 64, 64)  # Darker grey for better header readability
COLOUR_LIGHT_GREY = RGBColor(242, 242, 242)
COLOUR_CREAM = RGBColor(255, 250, 240)
COLOUR_YELLOW = RGBColor(255, 255, 0)
COLOUR_WHITE = RGBColor(255, 255, 255)
COLOUR_BLUE = RGBColor(68, 114, 196)  # For backwards compatibility with existing tables

# Sizes
SIZE_TITLE = Pt(20)  # Updated from 18pt per user preference
SIZE_SUBTITLE = Pt(14)
SIZE_HEADING2 = Pt(14)
SIZE_HEADING3 = Pt(12)
SIZE_BODY = Pt(12)
SIZE_CAPTION = Pt(10)
SIZE_HEADER_FOOTER = Pt(10)  # Updated from 9pt per user preference
SIZE_QUESTION = Pt(11)  # Updated from 12pt to match spec (Aptos, 11pt, Bold)

# Page Layout - Margins (Top: 1cm, Bottom: 1cm, Left/Right: 2cm)
MARGIN_TOP = Cm(1)
MARGIN_BOTTOM = Cm(1)
MARGIN_LEFT = Cm(2)
MARGIN_RIGHT = Cm(2)
MARGIN_DEFAULT = Cm(2)  # For backwards compatibility
PAGE_WIDTH_A4 = Cm(21)
PAGE_HEIGHT_A4 = Cm(29.7)
CONTENT_WIDTH = Cm(17)  # A4 (21cm) minus left (2cm) and right (2cm) margins

# Header/Footer heights
HEADER_HEIGHT = Cm(0.8)  # Updated from 0.6cm per user preference
FOOTER_HEIGHT = Cm(0.8)  # Updated from 1.25cm to match header height

# Usable page height calculation (for extended responses)
USABLE_HEIGHT = PAGE_HEIGHT_A4 - MARGIN_TOP - MARGIN_BOTTOM - HEADER_HEIGHT - FOOTER_HEIGHT  # ~27.25cm

# Spacing
SPACING_AFTER_PARA = Pt(6)
SPACING_BEFORE_HEADING = Pt(12)
SPACING_AFTER_HEADING = Pt(6)
LINE_SPACING = 1.15

# Blank Lines
BLANK_LINE_HEIGHT = Cm(0.7)


# =============================================================================
# COLOUR SCHEMES - Selectable palettes for document styling
# =============================================================================

COLOUR_SCHEMES = {
    'professional_minimal': {
        'heading': RGBColor(0, 0, 0),           # Black
        'body': RGBColor(0, 0, 0),              # Black
        'header_footer': RGBColor(64, 64, 64),  # Dark Grey
        'table_header_bg': RGBColor(242, 242, 242),  # Light Grey
        'quote_bg': RGBColor(255, 250, 240),    # Cream
        'quote_border': RGBColor(64, 64, 64),   # Dark Grey
        'instruction_bg': RGBColor(242, 242, 242),   # Light Grey
        'instruction_border': RGBColor(0, 0, 0),    # Black
        'highlight_evidence': RGBColor(255, 255, 0),    # Yellow
        'highlight_terms': RGBColor(242, 242, 242),     # Light Grey
        'table_border': RGBColor(0, 0, 0),     # Black
    },
    'academic_blue': {
        'heading': RGBColor(25, 55, 109),       # Deep Blue
        'body': RGBColor(40, 40, 40),           # Charcoal
        'header_footer': RGBColor(31, 78, 121),  # Navy Blue
        'table_header_bg': RGBColor(217, 225, 242),   # Light Blue
        'quote_bg': RGBColor(239, 243, 249),    # Pale Blue
        'quote_border': RGBColor(25, 55, 109),  # Deep Blue
        'instruction_bg': RGBColor(217, 225, 242),   # Light Blue
        'instruction_border': RGBColor(25, 55, 109),  # Deep Blue
        'highlight_evidence': RGBColor(255, 192, 0),   # Amber Yellow
        'highlight_terms': RGBColor(217, 225, 242),    # Light Blue
        'table_border': RGBColor(31, 78, 121),   # Navy Blue
    },
    'nature_green': {
        'heading': RGBColor(34, 94, 56),        # Forest Green
        'body': RGBColor(40, 40, 40),           # Charcoal
        'header_footer': RGBColor(114, 145, 110),  # Sage Green
        'table_header_bg': RGBColor(220, 237, 220),   # Mint Green
        'quote_bg': RGBColor(242, 250, 242),    # Pale Mint
        'quote_border': RGBColor(34, 94, 56),   # Forest Green
        'instruction_bg': RGBColor(220, 237, 220),   # Mint Green
        'instruction_border': RGBColor(34, 94, 56),  # Forest Green
        'highlight_evidence': RGBColor(218, 165, 32),   # Warm Gold
        'highlight_terms': RGBColor(220, 237, 220),     # Mint Green
        'table_border': RGBColor(114, 145, 110),   # Sage Green
    },
    'creative_vibrant': {
        'heading': RGBColor(75, 0, 130),        # Deep Purple
        'body': RGBColor(40, 40, 40),           # Charcoal
        'header_footer': RGBColor(147, 51, 234),  # Medium Purple
        'table_header_bg': RGBColor(230, 220, 245),   # Light Lavender
        'quote_bg': RGBColor(245, 240, 250),    # Pale Lavender
        'quote_border': RGBColor(75, 0, 130),   # Deep Purple
        'instruction_bg': RGBColor(230, 220, 245),   # Light Lavender
        'instruction_border': RGBColor(75, 0, 130),  # Deep Purple
        'highlight_evidence': RGBColor(255, 140, 0),   # Vibrant Orange
        'highlight_terms': RGBColor(230, 220, 245),    # Light Lavender
        'table_border': RGBColor(147, 51, 234),   # Medium Purple
    },
    'warm_humanities': {
        'heading': RGBColor(154, 48, 19),       # Rust Red
        'body': RGBColor(40, 40, 40),           # Charcoal
        'header_footer': RGBColor(191, 87, 0),   # Burnt Sienna
        'table_header_bg': RGBColor(242, 220, 200),   # Pale Terracotta
        'quote_bg': RGBColor(255, 250, 240),    # Cream
        'quote_border': RGBColor(154, 48, 19),  # Rust Red
        'instruction_bg': RGBColor(242, 220, 200),   # Pale Terracotta
        'instruction_border': RGBColor(154, 48, 19),  # Rust Red
        'highlight_evidence': RGBColor(184, 134, 11),   # Deep Gold
        'highlight_terms': RGBColor(242, 220, 200),    # Pale Terracotta
        'table_border': RGBColor(191, 87, 0),    # Burnt Sienna
    },
}

# Default colour scheme
DEFAULT_COLOUR_SCHEME = 'professional_minimal'


# =============================================================================
# STYLE NAMES - Using Word's built-in styles (modified at document creation)
# =============================================================================

# These are Word's built-in style names - we modify them at document creation
# to match our spec. Changes are document-scoped (don't affect Normal.dotm).
STYLE_TITLE = 'Title'                 # Document title (20pt, centred)
STYLE_SUBTITLE = 'Subtitle'           # Subtitle (14pt, centred)
STYLE_HEADING1 = 'Heading 1'          # Main heading (available if needed)
STYLE_HEADING2 = 'Heading 2'          # Section heading (14pt, with border)
STYLE_HEADING3 = 'Heading 3'          # Subsection heading (12pt)
STYLE_BODY = 'Normal'                 # Normal body text (12pt)
STYLE_QUESTION = 'Normal'             # Questions use Normal (styled inline)
STYLE_INSTRUCTION = 'Normal'          # Instructions use Normal (styled inline)
STYLE_QUOTE = 'Quote'                 # Quote/extract (12pt, italic)
STYLE_CAPTION = 'Caption'             # Image caption (10pt, italic, grey)


# =============================================================================
# DOCUMENT SETUP WITH STYLES
# =============================================================================

def setup_document(colour_scheme='professional_minimal'):
    """
    Create a new document with standard page setup and modified built-in styles.

    Styles modified (document-scoped, does not affect Word's global defaults):
    - Title: Document title (20pt Aptos Display, Bold, Centre)
    - Subtitle: Subtitle (14pt Aptos, Bold, Centre)
    - Heading 1: Main heading (16pt Aptos Display, Bold)
    - Heading 2: Section heading (14pt Aptos Display, Bold)
    - Heading 3: Subsection heading (12pt Aptos, Bold)
    - Normal: Body text (12pt Aptos)
    - Quote: Quote text (12pt Aptos, Italic)
    - Caption: Image caption (10pt Aptos, Italic, Grey)

    Args:
        colour_scheme: Colour palette to use. Options:
            - 'professional_minimal' (default) - formal, B&W print-friendly
            - 'academic_blue' - professional blue theme
            - 'nature_green' - green/natural theme
            - 'creative_vibrant' - purple/creative theme
            - 'warm_humanities' - warm red/brown theme

    Returns:
        Document: Configured document with A4 page size, custom margins
                  (Top: 1cm, Bottom/Left/Right: 2cm) and selected colour scheme
    """
    doc = Document()

    # Validate colour scheme
    if colour_scheme not in COLOUR_SCHEMES:
        raise ValueError(
            f"Unknown colour scheme: {colour_scheme}. "
            f"Available schemes: {', '.join(COLOUR_SCHEMES.keys())}"
        )

    # Store colour scheme in document object for use by other functions
    doc.colour_scheme = colour_scheme
    doc.colours = COLOUR_SCHEMES[colour_scheme]

    # Configure page layout (first section)
    section = doc.sections[0]
    section.page_width = PAGE_WIDTH_A4
    section.page_height = PAGE_HEIGHT_A4
    section.orientation = WD_ORIENT.PORTRAIT
    section.left_margin = MARGIN_LEFT
    section.right_margin = MARGIN_RIGHT
    section.top_margin = MARGIN_TOP
    section.bottom_margin = MARGIN_BOTTOM
    section.header_distance = HEADER_HEIGHT  # Distance from top of page to header

    # Modify built-in styles for this document
    _create_styles(doc)

    return doc


def _create_styles(doc):
    """
    Modify Word's built-in styles for this document.

    These modifications are document-scoped and do NOT affect:
    - Word's global Normal.dotm template
    - Other documents
    - The user's default styles

    This keeps the Style Gallery clean while ensuring consistent formatting.
    """
    styles = doc.styles

    # Get colours from the document's selected scheme
    colours = doc.colours

    # --- Normal Style (Body Text) ---
    # This is the base style - must be configured first
    normal_style = styles['Normal']
    normal_style.font.name = FONT_PRIMARY
    normal_style.font.size = SIZE_BODY
    normal_style.font.color.rgb = colours['body']
    normal_style.paragraph_format.space_after = SPACING_AFTER_PARA
    normal_style.paragraph_format.line_spacing = LINE_SPACING

    # --- Title Style ---
    title_style = styles['Title']
    title_style.font.name = FONT_DISPLAY
    title_style.font.size = SIZE_TITLE  # 20pt
    title_style.font.bold = True
    title_style.font.color.rgb = colours['heading']
    title_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_style.paragraph_format.space_before = Pt(12)  # Space from header
    title_style.paragraph_format.space_after = Pt(8)  # Space before next element
    _remove_style_borders(title_style)
    _set_style_font_xml(title_style, FONT_DISPLAY)  # Ensure font at XML level

    # --- Subtitle Style ---
    subtitle_style = styles['Subtitle']
    subtitle_style.font.name = FONT_PRIMARY
    subtitle_style.font.size = SIZE_SUBTITLE
    subtitle_style.font.bold = True
    subtitle_style.font.color.rgb = colours['heading']
    subtitle_style.font.italic = False  # Override default italic
    subtitle_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_style.paragraph_format.space_before = Pt(0)
    subtitle_style.paragraph_format.space_after = Pt(12)
    _remove_style_borders(subtitle_style)
    _set_style_font_xml(subtitle_style, FONT_PRIMARY)  # Ensure font at XML level

    # --- Heading 1 Style ---
    h1_style = styles['Heading 1']
    h1_style.font.name = FONT_DISPLAY
    h1_style.font.size = Pt(18)  # Main section headings
    h1_style.font.bold = True
    h1_style.font.color.rgb = colours['heading']
    h1_style.font.underline = False  # Remove default underline
    h1_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    h1_style.paragraph_format.space_before = SPACING_BEFORE_HEADING
    h1_style.paragraph_format.space_after = SPACING_AFTER_HEADING
    _set_style_font_xml(h1_style, FONT_DISPLAY)  # Ensure font at XML level

    # --- Heading 2 Style ---
    h2_style = styles['Heading 2']
    h2_style.font.name = FONT_DISPLAY
    h2_style.font.size = Pt(16)  # Subsection headings
    h2_style.font.bold = True
    h2_style.font.color.rgb = colours['heading']
    h2_style.font.underline = False  # Remove default underline
    h2_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    h2_style.paragraph_format.space_before = SPACING_BEFORE_HEADING
    h2_style.paragraph_format.space_after = SPACING_AFTER_HEADING
    _set_style_font_xml(h2_style, FONT_DISPLAY)  # Ensure font at XML level

    # --- Heading 3 Style ---
    h3_style = styles['Heading 3']
    h3_style.font.name = FONT_PRIMARY
    h3_style.font.size = SIZE_HEADING3
    h3_style.font.bold = True
    h3_style.font.color.rgb = colours['heading']
    h3_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    h3_style.paragraph_format.space_before = Pt(10)
    h3_style.paragraph_format.space_after = Pt(4)
    _set_style_font_xml(h3_style, FONT_PRIMARY)  # Ensure font at XML level

    # --- Quote Style ---
    quote_style = styles['Quote']
    quote_style.font.name = FONT_PRIMARY
    quote_style.font.size = SIZE_BODY
    quote_style.font.italic = True
    quote_style.font.color.rgb = COLOUR_BLACK
    quote_style.paragraph_format.line_spacing = 1.5
    quote_style.paragraph_format.space_after = SPACING_AFTER_PARA

    # --- Caption Style ---
    caption_style = styles['Caption']
    caption_style.font.name = FONT_PRIMARY
    caption_style.font.size = SIZE_CAPTION
    caption_style.font.italic = True
    caption_style.font.color.rgb = COLOUR_GREY
    caption_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption_style.paragraph_format.space_after = Pt(12)


def add_header_footer(doc, year_level, unit_name, doc_type=None, include_name=True):
    """
    Add header (year level, unit name, document type, and optional name field) and footer (page numbers).

    The header uses a table layout:
    - Left: "Year X English - Unit Name - Doc Type"
    - Right: "Name: ____________" (if include_name=True, first page only)

    Args:
        doc: Document object
        year_level: e.g., "10" or "Year 10"
        unit_name: e.g., "Media and Advertising"
        doc_type: Document type - one of: "Worksheet", "Handout", "Quiz", "Model Answer"
                  If None, omitted from header
        include_name: Whether to include Name field in first-page header (default True)
    """
    section = doc.sections[0]

    # Enable different first page header/footer
    section.different_first_page_header_footer = True

    # Ensure year_level is formatted correctly
    if not str(year_level).lower().startswith('year'):
        year_level = f"Year {year_level}"

    # Build header text
    header_text = f"{year_level} English - {unit_name}"
    if doc_type:
        header_text += f" - {doc_type}"

    # Header with table layout for left/right alignment (first page)
    header = section.first_page_header

    # Clear existing content
    for para in header.paragraphs:
        para.clear()

    if include_name:
        # Use a 2-column table for layout: left text, right name field
        header_table = header.add_table(rows=1, cols=2, width=CONTENT_WIDTH)
        header_table.autofit = False
        header_table.allow_autofit = False

        # Set table width to full content width
        tbl = header_table._tbl
        tblPr = tbl.tblPr if tbl.tblPr is not None else OxmlElement('w:tblPr')
        tblW = OxmlElement('w:tblW')
        tblW.set(qn('w:w'), str(int(CONTENT_WIDTH.twips)))
        tblW.set(qn('w:type'), 'dxa')
        tblPr.append(tblW)

        # Remove table borders
        tblBorders = OxmlElement('w:tblBorders')
        for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
            border = OxmlElement(f'w:{border_name}')
            border.set(qn('w:val'), 'nil')
            tblBorders.append(border)
        tblPr.append(tblBorders)

        row = header_table.rows[0]
        left_cell = row.cells[0]
        right_cell = row.cells[1]

        # Set column widths (left: ~60%, right: ~40% for extended name line)
        left_cell.width = Cm(10)
        right_cell.width = Cm(6)

        # Left cell: document info
        left_para = left_cell.paragraphs[0]
        left_run = left_para.add_run(header_text)
        left_run.font.name = FONT_PRIMARY
        left_run.font.size = SIZE_HEADER_FOOTER
        left_run.font.color.rgb = COLOUR_DARK_GREY  # Darker grey for better readability
        left_para.alignment = WD_ALIGN_PARAGRAPH.LEFT

        # Right cell: Name field with underscores
        right_para = right_cell.paragraphs[0]
        name_run = right_para.add_run("Name: ")
        name_run.font.name = FONT_PRIMARY
        name_run.font.size = SIZE_HEADER_FOOTER
        name_run.font.color.rgb = COLOUR_DARK_GREY  # Darker grey for better readability

        # Add underscores for the line (26 chars = 30% longer for extended name space)
        line_run = right_para.add_run("_" * 26)
        line_run.font.name = FONT_PRIMARY
        line_run.font.size = SIZE_HEADER_FOOTER
        line_run.font.color.rgb = COLOUR_DARK_GREY  # Darker grey for better readability
        right_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    else:
        # Simple header without name field (original behaviour)
        header_para = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
        header_para.clear()
        header_run = header_para.add_run(header_text)
        header_run.font.name = FONT_PRIMARY
        header_run.font.size = SIZE_HEADER_FOOTER
        header_run.font.color.rgb = COLOUR_DARK_GREY  # Darker grey for better readability
        header_para.alignment = WD_ALIGN_PARAGRAPH.LEFT

    # Set up default header for subsequent pages (without name field)
    default_header = section.header
    for para in default_header.paragraphs:
        para.clear()

    default_header_para = default_header.paragraphs[0] if default_header.paragraphs else default_header.add_paragraph()
    default_header_para.clear()
    default_header_run = default_header_para.add_run(header_text)
    default_header_run.font.name = FONT_PRIMARY
    default_header_run.font.size = SIZE_HEADER_FOOTER
    default_header_run.font.color.rgb = COLOUR_DARK_GREY
    default_header_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    default_header_para.paragraph_format.space_after = Pt(6)  # Add spacing to match table-based header

    # Footer with page number
    footer = section.footer
    footer_para = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    footer_para.clear()
    footer_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    # Add page number field (must be wrapped in runs for valid OOXML)
    run1 = OxmlElement('w:r')
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    run1.append(fldChar1)
    footer_para._p.append(run1)

    run2 = OxmlElement('w:r')
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = " PAGE "
    run2.append(instrText)
    footer_para._p.append(run2)

    run3 = OxmlElement('w:r')
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'end')
    run3.append(fldChar2)
    footer_para._p.append(run3)


# =============================================================================
# TITLES AND HEADINGS (Using Styles)
# =============================================================================

def add_title(doc, text, centre=True):
    """
    Add document title using EA Title style.

    Args:
        doc: Document object
        text: Title text
        centre: Whether to centre the title (default True)

    Returns:
        Paragraph object
    """
    para = doc.add_paragraph(text, style=STYLE_TITLE)
    if not centre:
        para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    return para


def add_subtitle(doc, text, centre=True):
    """
    Add subtitle using EA Subtitle style.

    Args:
        doc: Document object
        text: Subtitle text
        centre: Whether to centre (default True)

    Returns:
        Paragraph object
    """
    para = doc.add_paragraph(text, style=STYLE_SUBTITLE)
    if not centre:
        para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    return para


def add_section_heading(doc, text):
    """
    Add section heading using Heading 1 style.

    Args:
        doc: Document object
        text: Heading text

    Returns:
        Paragraph object
    """
    return doc.add_paragraph(text, style=STYLE_HEADING1)


def add_subsection_heading(doc, text):
    """
    Add subsection heading using Heading 2 style.

    Args:
        doc: Document object
        text: Heading text

    Returns:
        Paragraph object
    """
    return doc.add_paragraph(text, style=STYLE_HEADING2)


# =============================================================================
# STUDENT INFORMATION
# =============================================================================

def add_name_date_block(doc, include_date=False):
    """
    Add Name block for student worksheets.
    Creates a borderless table with underlined input area.

    Args:
        doc: Document object
        include_date: Whether to include date field (default False for cleaner appearance)

    Returns:
        Table object
    """
    if include_date:
        # Legacy 4-column layout with date
        table = doc.add_table(rows=1, cols=4)
        table.autofit = False
        table.alignment = WD_TABLE_ALIGNMENT.LEFT

        row = table.rows[0]
        cells = row.cells

        # Name label
        name_label = cells[0].paragraphs[0]
        name_run = name_label.add_run("Name:")
        name_run.bold = True
        name_run.font.name = FONT_PRIMARY
        name_run.font.size = SIZE_BODY
        cells[0].width = Cm(1.5)

        # Name input (underlined space)
        cells[1].width = Cm(8)
        _add_underlined_cell(cells[1])

        # Date label
        date_label = cells[2].paragraphs[0]
        date_run = date_label.add_run("Date:")
        date_run.bold = True
        date_run.font.name = FONT_PRIMARY
        date_run.font.size = SIZE_BODY
        cells[2].width = Cm(1.5)

        # Date input (underlined space)
        cells[3].width = Cm(4)
        _add_underlined_cell(cells[3])
    else:
        # New simplified 2-column layout (name only, 60% width)
        table = doc.add_table(rows=1, cols=2)
        table.autofit = False
        table.alignment = WD_TABLE_ALIGNMENT.LEFT

        row = table.rows[0]
        cells = row.cells

        # Name label
        name_label = cells[0].paragraphs[0]
        name_run = name_label.add_run("Name:")
        name_run.bold = True
        name_run.font.name = FONT_PRIMARY
        name_run.font.size = SIZE_BODY
        cells[0].width = Cm(1.5)

        # Name input (underlined space - ~60% of previous 8cm = ~5cm)
        cells[1].width = Cm(5)
        _add_underlined_cell(cells[1])

    # Remove all borders from table
    _remove_table_borders(table)

    # Add spacing after
    doc.add_paragraph()

    return table


def add_name_block(doc):
    """
    Alias for add_name_date_block without date.
    Simplified name-only block for cleaner worksheets.

    Args:
        doc: Document object

    Returns:
        Table object
    """
    return add_name_date_block(doc, include_date=False)


def _add_underlined_cell(cell):
    """Add bottom border only to a cell to create underlined input area."""
    tcPr = cell._tc.get_or_add_tcPr()
    tcBorders = parse_xml(
        r'<w:tcBorders xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
        r'<w:top w:val="nil"/>'
        r'<w:left w:val="nil"/>'
        r'<w:right w:val="nil"/>'
        r'<w:bottom w:val="single" w:sz="4" w:color="000000"/>'
        r'</w:tcBorders>'
    )
    tcPr.append(tcBorders)


def _remove_table_borders(table):
    """Remove all borders from a table."""
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else OxmlElement('w:tblPr')
    tblBorders = OxmlElement('w:tblBorders')
    for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'nil')
        tblBorders.append(border)
    tblPr.append(tblBorders)
    if tbl.tblPr is None:
        tbl.insert(0, tblPr)


# =============================================================================
# BLANK LINES FOR RESPONSES
# =============================================================================

def add_blank_lines(doc, num_lines=2, add_spacing=True):
    """
    Add blank lines for student responses using table method.

    Args:
        doc: Document object
        num_lines: Number of lines (default 2)
            - 2 lines: Short response (1-2 sentences)
            - 4 lines: Medium response (3-4 sentences)
            - 6 lines: Extended response (paragraph)
            - 8+ lines: Long response
        add_spacing: Whether to add spacing paragraph after table (default True)

    Returns:
        Table object
    """
    table = doc.add_table(rows=num_lines, cols=1)
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT

    # Set table width to content width
    table.columns[0].width = CONTENT_WIDTH

    for idx, row in enumerate(table.rows):
        # First row is smaller to reduce space after question
        if idx == 0:
            row.height = Cm(0.35)  # Half height for first row
        else:
            row.height = BLANK_LINE_HEIGHT
        row.height_rule = WD_ROW_HEIGHT_RULE.EXACTLY

        cell = row.cells[0]
        # Remove all borders except bottom
        tcPr = cell._tc.get_or_add_tcPr()
        tcBorders = parse_xml(
            r'<w:tcBorders xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
            r'<w:top w:val="nil"/>'
            r'<w:left w:val="nil"/>'
            r'<w:right w:val="nil"/>'
            r'<w:bottom w:val="single" w:sz="4" w:color="000000"/>'
            r'</w:tcBorders>'
        )
        tcPr.append(tcBorders)

        # Remove cell padding
        tcMar = OxmlElement('w:tcMar')
        for margin_name in ['top', 'left', 'bottom', 'right']:
            margin = OxmlElement(f'w:{margin_name}')
            margin.set(qn('w:w'), '0')
            margin.set(qn('w:type'), 'dxa')
            tcMar.append(margin)
        tcPr.append(tcMar)

    # Add spacing after blank lines table (unless it's a section-ending response)
    if add_spacing:
        spacing_para = doc.add_paragraph()
        spacing_para.paragraph_format.space_before = Pt(0)
        spacing_para.paragraph_format.space_after = Pt(6)

    return table


def add_half_page_response(doc):
    """
    Add a half-page response area for extended writing.
    Calculates line count based on page margins to fill approximately half the usable page.

    Args:
        doc: Document object

    Returns:
        Table object
    """
    # Calculate lines for half page
    # Usable height ~24.85cm, half = ~12.4cm, each line = 0.7cm
    # 12.4 / 0.7 ≈ 18 lines
    num_lines = 18
    return add_blank_lines(doc, num_lines, add_spacing=False)


def add_full_page_response(doc):
    """
    Add a full-page response area for extended writing.
    Calculates line count based on page margins to fill the entire usable page.

    Args:
        doc: Document object

    Returns:
        Table object
    """
    # Calculate lines for full page
    # Usable height = 29.7 - 1 (top) - 1 (bottom) - 0.8 (header) - 0.8 (footer) = ~27.25cm
    # First line = 0.35cm, remaining lines = 0.7cm each
    # Account for section heading and body text (~1.5cm), leaves ~25.75cm
    # (25.75 - 0.35) / 0.7 ≈ 36 lines, use 33 to prevent overflow onto next page
    num_lines = 33
    return add_blank_lines(doc, num_lines, add_spacing=False)


def add_extended_response(doc, size='half'):
    """
    Add an extended writing response area.

    Args:
        doc: Document object
        size: 'half' for half-page (~17 lines) or 'full' for full-page (~33 lines)

    Returns:
        Table object
    """
    if size == 'full':
        return add_full_page_response(doc)
    else:
        return add_half_page_response(doc)


# =============================================================================
# TABLES
# =============================================================================

def add_analysis_table(doc, data, has_header_bg=True, use_blue=False):
    """
    Create a formatted analysis table (Feature/Example/Effect/Purpose style).

    Args:
        doc: Document object
        data: List of tuples [(key, value), ...] or dict {key: value}
        has_header_bg: Whether to apply background to first column (default True)
        use_blue: Use blue background (legacy) or grey (spec default)

    Returns:
        Table object
    """
    if isinstance(data, dict):
        data = list(data.items())

    table = doc.add_table(rows=len(data), cols=2)
    # Calculate optimal left column width based on longest label
    # Heuristic: 0.23 cm per character + 0.5 cm padding (V15)
    # V14 (2.06cm) was too narrow and wrapped. 12pt Bold needs more space.
    # "Connotations" (12 chars) -> 12 * 0.23 + 0.5 = 3.26 cm
    max_label_len = 0
    for key, val in data:
        max_label_len = max(max_label_len, len(str(key)))

    # "Connotations" (12 chars) -> 12 * 0.23 + 0.5 = 3.26 cm
    optimal_width = (max_label_len * 0.23) + 0.5

    # Enforce a reasonable minimum (1.0 cm) and maximum (3.5 cm)
    optimal_width = max(1.0, min(optimal_width, 3.5))

    table.style = 'Table Grid'
    table.autofit = False  # Critical: Disable autofit to enforce our calculated width
    table.allow_autofit = False

    table.columns[0].width = Cm(optimal_width)
    table.columns[1].width = Cm(17.0 - optimal_width)  # Remainder of 17cm page width

    for i, (key, value) in enumerate(data):
        row = table.rows[i]
        row.height_rule = WD_ROW_HEIGHT_RULE.AUTO  # Let content determine height

        # Set vertical alignment and padding for both cells in the row
        for cell in row.cells:
            _set_vertical_alignment(cell, 'center')
            _set_cell_padding(cell, Cm(0.19))  # Consistent padding: 0.19cm

            # Remove line/paragraph spacing inside table cells for perfect centering
            for para in cell.paragraphs:
                para.paragraph_format.space_before = Pt(0)
                para.paragraph_format.space_after = Pt(0)
                para.paragraph_format.line_spacing = 1.0  # Force single spacing

        # First cell (key) - Labels are Bold and Left-Aligned
        key_cell = row.cells[0]
        key_para = key_cell.paragraphs[0]
        key_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        key_run = key_para.add_run(str(key))
        key_run.bold = True
        key_run.font.name = FONT_PRIMARY
        key_run.font.size = SIZE_BODY

        # Apply background to header column if requested
        if has_header_bg:
            shading_elm = OxmlElement('w:shd')
            if use_blue:
                shading_elm.set(qn('w:fill'), '4472C4')  # Blue (legacy)
                key_run.font.color.rgb = COLOUR_WHITE
            else:
                shading_elm.set(qn('w:fill'), 'F2F2F2')  # Light grey (spec)
                key_run.font.color.rgb = COLOUR_BLACK
            key_cell._element.get_or_add_tcPr().append(shading_elm)

        # Second cell (value)
        value_cell = row.cells[1]
        value_para = value_cell.paragraphs[0]
        value_run = value_para.add_run(str(value))
        value_run.font.name = FONT_PRIMARY
        value_run.font.size = SIZE_BODY

        # Explicitly set cell widths to ensure Word respects them
        row.cells[0].width = Cm(optimal_width)
        row.cells[1].width = Cm(17.0 - optimal_width)

    return table


def add_content_table(doc, headers, rows, header_bg=True):
    """
    Create a standard content table with header row.

    Args:
        doc: Document object
        headers: List of header strings
        rows: List of row data (each row is a list of cell values)
        header_bg: Whether to apply grey background to header row

    Returns:
        Table object
    """
    table = doc.add_table(rows=len(rows) + 1, cols=len(headers))
    table.style = 'Table Grid'

    # Header row
    header_row = table.rows[0]
    header_row.height_rule = WD_ROW_HEIGHT_RULE.AUTO  # Auto height for header
    for i, header_text in enumerate(headers):
        cell = header_row.cells[i]
        para = cell.paragraphs[0]
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER  # Centre align header text
        # Remove paragraph spacing for compact height
        para.paragraph_format.space_before = Pt(0)
        para.paragraph_format.space_after = Pt(0)
        para.paragraph_format.line_spacing = 1.0

        run = para.add_run(str(header_text))
        run.bold = True
        run.font.name = FONT_PRIMARY
        run.font.size = SIZE_BODY

        if header_bg:
            shading_elm = OxmlElement('w:shd')
            shading_elm.set(qn('w:fill'), 'F2F2F2')
            cell._element.get_or_add_tcPr().append(shading_elm)

        # Add consistent padding (0.19cm) and vertical centre alignment
        _set_cell_padding(cell, Cm(0.19))
        _set_vertical_alignment(cell, 'center')

    # Data rows
    for row_idx, row_data in enumerate(rows):
        table_row = table.rows[row_idx + 1]
        table_row.height_rule = WD_ROW_HEIGHT_RULE.AUTO  # Auto height for data rows

        for col_idx, cell_value in enumerate(row_data):
            cell = table_row.cells[col_idx]
            para = cell.paragraphs[0]
            # Remove paragraph spacing for compact height
            para.paragraph_format.space_before = Pt(0)
            para.paragraph_format.space_after = Pt(0)
            para.paragraph_format.line_spacing = 1.0

            run = para.add_run(str(cell_value))
            run.font.name = FONT_PRIMARY
            run.font.size = SIZE_BODY
            # Add consistent padding (0.19cm) and vertical centre alignment
            _set_cell_padding(cell, Cm(0.19))
            _set_vertical_alignment(cell, 'center')

    return table


# =============================================================================
# BOXES AND CALLOUTS
# =============================================================================

def add_instruction_box(doc, text):
    """
    Add an instruction box with grey background.

    Args:
        doc: Document object
        text: Instruction text

    Returns:
        Table object
    """
    # Create a single-cell table for the box
    table = doc.add_table(rows=1, cols=1)
    table.autofit = False
    table.columns[0].width = CONTENT_WIDTH

    cell = table.rows[0].cells[0]

    # Add shading
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), 'F2F2F2')
    cell._element.get_or_add_tcPr().append(shading_elm)

    # Add text with instruction style formatting
    para = cell.paragraphs[0]
    run = para.add_run(text)
    run.italic = True
    run.font.name = FONT_PRIMARY
    run.font.size = SIZE_BODY

    # Add border
    _set_cell_border(cell, '000000', '8')  # 1pt border

    # Add padding (0.19cm = Word default)
    _set_cell_padding(cell, Cm(0.19))

    # Add spacing after the box (reduced 30% from 6pt to 4pt)
    spacing_para = doc.add_paragraph()
    spacing_para.paragraph_format.space_before = Pt(0)
    spacing_para.paragraph_format.space_after = Pt(4)

    return table


def add_quote_box(doc, text, source=None):
    """
    Add a quote/extract box with left border and cream background.
    Text uses EA Quote style (italic).

    Args:
        doc: Document object
        text: Quote/extract text
        source: Optional source attribution

    Returns:
        Table object
    """
    # Create a single-cell table for the box
    table = doc.add_table(rows=1, cols=1)
    table.autofit = False
    table.columns[0].width = CONTENT_WIDTH

    cell = table.rows[0].cells[0]

    # Add cream background
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), 'FFFAF0')  # Light cream
    cell._element.get_or_add_tcPr().append(shading_elm)

    # Add text with quote style formatting
    para = cell.paragraphs[0]
    run = para.add_run(text)
    run.italic = True
    run.font.name = FONT_PRIMARY
    run.font.size = SIZE_BODY
    para.paragraph_format.line_spacing = 1.5

    # Add source if provided
    if source:
        source_para = cell.add_paragraph()
        source_run = source_para.add_run(f"— {source}")
        source_run.font.name = FONT_PRIMARY
        source_run.font.size = SIZE_CAPTION
        source_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    # Add left border only (2pt dark grey)
    tcPr = cell._tc.get_or_add_tcPr()
    tcBorders = parse_xml(
        r'<w:tcBorders xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
        r'<w:top w:val="nil"/>'
        r'<w:right w:val="nil"/>'
        r'<w:bottom w:val="nil"/>'
        r'<w:left w:val="single" w:sz="16" w:color="808080"/>'
        r'</w:tcBorders>'
    )
    tcPr.append(tcBorders)

    # Add padding
    _set_cell_padding(cell, Cm(0.5))

    return table


def add_answer_box(doc, text):
    """
    Add an answer box (for model answers) with thin border.

    Args:
        doc: Document object
        text: Answer text

    Returns:
        Table object
    """
    table = doc.add_table(rows=1, cols=1)
    table.autofit = False
    table.columns[0].width = CONTENT_WIDTH

    cell = table.rows[0].cells[0]

    # Add text
    para = cell.paragraphs[0]
    run = para.add_run(text)
    run.font.name = FONT_PRIMARY
    run.font.size = SIZE_BODY

    # Add border
    _set_cell_border(cell, '000000', '4')  # 0.5pt border

    # Add padding
    _set_cell_padding(cell, Cm(0.3))

    return table


def _set_cell_border(cell, colour, size):
    """Set all borders on a cell."""
    tcPr = cell._tc.get_or_add_tcPr()
    tcBorders = parse_xml(
        f'<w:tcBorders xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
        f'<w:top w:val="single" w:sz="{size}" w:color="{colour}"/>'
        f'<w:left w:val="single" w:sz="{size}" w:color="{colour}"/>'
        f'<w:bottom w:val="single" w:sz="{size}" w:color="{colour}"/>'
        f'<w:right w:val="single" w:sz="{size}" w:color="{colour}"/>'
        f'</w:tcBorders>'
    )
    tcPr.append(tcBorders)


def _set_cell_padding(cell, padding):
    """Set padding on all sides of a cell."""
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    padding_twips = int(padding.twips) if hasattr(padding, 'twips') else int(padding * 567)  # Convert cm to twips
    for margin_name in ['top', 'left', 'bottom', 'right']:
        margin = OxmlElement(f'w:{margin_name}')
        margin.set(qn('w:w'), str(padding_twips))
        margin.set(qn('w:type'), 'dxa')
        tcMar.append(margin)
    tcPr.append(tcMar)


# =============================================================================
# NUMBERED QUESTIONS
# =============================================================================

def add_question(doc, number, text, bold_number=True):
    """
    Add a numbered question using EA Question style.

    Args:
        doc: Document object
        number: Question number
        text: Question text
        bold_number: Whether to bold the number (default True)

    Returns:
        Paragraph object
    """
    para = doc.add_paragraph(style=STYLE_QUESTION)

    num_run = para.add_run(f"{number}. ")
    num_run.font.name = FONT_PRIMARY
    num_run.font.size = SIZE_QUESTION
    if bold_number:
        num_run.bold = True

    text_run = para.add_run(text)
    text_run.font.name = FONT_PRIMARY
    text_run.font.size = SIZE_QUESTION

    # Set uniform spacing
    para.paragraph_format.space_before = Pt(0)
    para.paragraph_format.space_after = Pt(6)

    return para


def add_sub_question(doc, letter, text):
    """
    Add a sub-question (a) b) c) style).

    Args:
        doc: Document object
        letter: Sub-question letter (e.g., 'a', 'b')
        text: Question text

    Returns:
        Paragraph object
    """
    para = doc.add_paragraph(style=STYLE_BODY)

    letter_run = para.add_run(f"{letter}) ")
    letter_run.font.name = FONT_PRIMARY
    letter_run.font.size = SIZE_BODY

    text_run = para.add_run(text)
    text_run.font.name = FONT_PRIMARY
    text_run.font.size = SIZE_BODY

    para.paragraph_format.left_indent = Cm(1.27)
    para.paragraph_format.space_before = Pt(0)
    para.paragraph_format.space_after = Pt(6)

    return para


# =============================================================================
# HIGHLIGHTING
# =============================================================================

def add_highlighted_text(doc, text, highlight_colour='yellow'):
    """
    Add text with highlighting.

    Args:
        doc: Document object
        text: Text to highlight
        highlight_colour: 'yellow' or 'grey'

    Returns:
        Paragraph object
    """
    para = doc.add_paragraph(style=STYLE_BODY)
    run = para.add_run(text)

    # Set highlight
    if highlight_colour == 'yellow':
        run.font.highlight_color = 7  # WD_COLOR_INDEX.YELLOW
    elif highlight_colour == 'grey':
        run.font.highlight_color = 15  # WD_COLOR_INDEX.GRAY_25

    return para


# =============================================================================
# PAGE ELEMENTS
# =============================================================================

def add_horizontal_rule(doc):
    """
    Add a horizontal rule (section divider).

    Args:
        doc: Document object

    Returns:
        Paragraph object
    """
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(6)
    para.paragraph_format.space_after = Pt(6)

    # Add bottom border to paragraph
    pPr = para._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '4')
    bottom.set(qn('w:color'), '000000')
    pBdr.append(bottom)
    pPr.append(pBdr)

    return para


def add_page_break(doc):
    """
    Add a page break.

    Args:
        doc: Document object
    """
    doc.add_page_break()


# =============================================================================
# CHECKLISTS
# =============================================================================

def add_checklist(doc, items):
    """
    Add a checklist with checkbox characters.

    Args:
        doc: Document object
        items: List of checklist item strings

    Returns:
        List of Paragraph objects
    """
    paragraphs = []
    for item in items:
        para = doc.add_paragraph(style=STYLE_BODY)
        run = para.add_run(f"\u2610 {item}")  # Unicode checkbox
        para.paragraph_format.line_spacing = 1.5
        paragraphs.append(para)

    return paragraphs


# =============================================================================
# UTILITY FUNCTIONS
# =============================================================================

def add_body_paragraph(doc, text, bold=False, italic=False):
    """
    Add a standard body paragraph using EA Body style.

    Args:
        doc: Document object
        text: Paragraph text
        bold: Whether to bold the text
        italic: Whether to italicise the text

    Returns:
        Paragraph object
    """
    para = doc.add_paragraph(style=STYLE_BODY)
    run = para.add_run(text)
    run.bold = bold
    run.italic = italic

    return para


def add_caption(doc, text, figure_number=None):
    """
    Add an image caption using EA Caption style.

    Args:
        doc: Document object
        text: Caption text
        figure_number: Optional figure number

    Returns:
        Paragraph object
    """
    if figure_number:
        caption_text = f"Figure {figure_number}: {text}"
    else:
        caption_text = text

    return doc.add_paragraph(caption_text, style=STYLE_CAPTION)


def _set_vertical_alignment(cell, align='center'):
    """
    Set vertical alignment for a cell.
    align: 'top', 'center', 'bottom'
    """
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()

    # Remove any existing vAlign
    for child in tcPr.xpath('w:vAlign'):
        tcPr.remove(child)

    vAlign = OxmlElement('w:vAlign')
    vAlign.set(qn('w:val'), align)
    tcPr.append(vAlign)


def _set_cell_margins(cell, top=None, start=None, bottom=None, end=None):
    """
    Set internal padding for a cell (in twips).
    1440 twips = 1 inch
    108 twips = ~0.19 cm
    """
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()

    # Remove any existing tcMar
    for child in tcPr.xpath('w:tcMar'):
        tcPr.remove(child)

    tcMar = OxmlElement('w:tcMar')

    for side, value in [('top', top), ('left', start), ('bottom', bottom), ('right', end)]:
        if value is not None:
            node = OxmlElement(f'w:{side}')
            node.set(qn('w:w'), str(value))
            node.set(qn('w:type'), 'dxa')
            tcMar.append(node)

    tcPr.append(tcMar)


def _remove_style_borders(style):
    """
    Remove all borders from a style using XML.
    Necessary for some Word versions that default to having borders (like the blue line).
    """
    pPr = style.element.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    for bdr in ['top', 'left', 'bottom', 'right', 'between', 'bar']:
        el = OxmlElement(f'w:{bdr}')
        el.set(qn('w:val'), 'nil')
        pBdr.append(el)
    pPr.append(pBdr)


def _set_style_font_xml(style, font_name):
    """
    Ensure font is set at XML level for the style.
    This ensures Word properly applies the font even if the display name differs.
    """
    rPr = style.element.get_or_add_rPr()

    # Remove existing font settings
    for rFont in rPr.findall(qn('w:rFonts')):
        rPr.remove(rFont)

    # Add new font settings for all font types (ASCII, HAnsi, EastAsia, CS)
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:ascii'), font_name)
    rFonts.set(qn('w:hAnsi'), font_name)
    rFonts.set(qn('w:eastAsia'), font_name)
    rFonts.set(qn('w:cs'), font_name)
    rPr.insert(0, rFonts)


def save_document(doc, filepath):
    """
    Save document to file.

    Args:
        doc: Document object
        filepath: Output filepath
    """
    doc.save(filepath)
    print(f"[SUCCESS] Document saved: {filepath}")


# =============================================================================
# STYLE REFERENCE
# =============================================================================
"""
BUILT-IN STYLES MODIFIED BY THIS MODULE
=======================================

This module modifies Word's built-in styles at document creation time.
Changes are document-scoped and do NOT affect Word's global defaults.

| Style Name   | Element Type        | Font          | Size | Weight  | Other           |
|--------------|---------------------|---------------|------|---------|-----------------|
| Title        | Document title      | Aptos Display | 20pt | Bold    | Centre          |
| Subtitle     | Subtitle            | Aptos         | 14pt | Bold    | Centre          |
| Heading 1    | Main heading        | Aptos Display | 16pt | Bold    | Left            |
| Heading 2    | Section heading     | Aptos Display | 14pt | Bold    | Left            |
| Heading 3    | Subsection heading  | Aptos         | 12pt | Bold    | Left            |
| Normal       | Body text           | Aptos         | 12pt | Regular | Left            |
| Quote        | Quote/extract       | Aptos         | 12pt | Italic  | Line spacing 1.5|
| Caption      | Image caption       | Aptos         | 10pt | Italic  | Centre, grey    |

WHY BUILT-IN STYLES?
====================

Using built-in styles instead of custom styles:
- Keeps the Style Gallery clean (no extra "EA" styles)
- Works seamlessly with Word's formatting features
- Document-scoped changes don't affect your Normal.dotm template
- Teachers can easily modify styles in the generated document

HOW TO MODIFY STYLES IN WORD
============================

1. Open a generated document in Word
2. In the Styles Gallery (Home tab), find the style you want to modify
3. Right-click the style → "Modify..."
4. Change any formatting (font, size, colour, spacing, etc.)
5. Click OK - all elements using that style update automatically

Note: These changes only affect the current document unless you explicitly
check "New documents based on this template" in the Modify Style dialog.
"""
